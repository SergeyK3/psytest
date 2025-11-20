"""
Расширенный модуль для создания отчётов с поддержкой AI интерпретаций
"""
import os
from pathlib import Path
from typing import Dict, Optional
from docx import Document
from docx.shared import Pt
import pandas as pd
from datetime import datetime

from .report import render_report as _render_report_base, _pick_interpretation
from .ai_interpreter import get_ai_interpreter


def render_enhanced_report(
    scores_raw: Dict[str, float],
    items_df: pd.DataFrame,
    interpretations_path: Path,
    template_path: Path,
    out_path: Path,
    title: str = "Отчёт по психологическому тестированию",
    participant: str = "Код участника: ______",
    test_type: str = "PAEI",
    use_ai: bool = None,
    dialog_context: str = ""
) -> Path:
    """
    Расширенная функция создания отчёта с поддержкой AI интерпретаций
    
    Args:
        scores_raw: Словарь сырых баллов
        items_df: DataFrame с вопросами
        interpretations_path: Путь к CSV с интерпретациями
        template_path: Путь к шаблону документа
        out_path: Путь для сохранения отчёта
        title: Заголовок отчёта
        participant: Информация об участнике
        test_type: Тип теста (PAEI, DISC, HEXACO)
        use_ai: Использовать ли AI (None = автоопределение)
        dialog_context: Контекст диалога для AI
        
    Returns:
        Путь к созданному отчёту
    """
    
    # Определяем, использовать ли AI
    if use_ai is None:
        use_ai = os.getenv("USE_AI_INTERPRETATIONS", "false").lower() == "true"
    
    # Пытаемся получить AI интерпретатор
    ai_interpreter = None
    if use_ai:
        ai_interpreter = get_ai_interpreter()
        if ai_interpreter is None:
            use_ai = False  # Fallback к статическим интерпретациям
    
    # Создаём базовый документ
    doc = Document(str(template_path))
    
    # Базовая типографика
    styles = doc.styles
    styles['Normal'].font.name = 'Times New Roman'
    styles['Normal'].font.size = Pt(12)
    
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"{participant}")
    doc.add_paragraph(f"Дата: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    # Подготовим нормировку к 0..60
    items_per_scale = items_df.groupby("scale")["item_id"].count().to_dict()
    max_per_scale = {k: v * 5 for k, v in items_per_scale.items()}
    scaled_scores = {}
    for s, raw in scores_raw.items():
        max_possible = max_per_scale.get(s, max(raw, 1))
        scaled_scores[s] = round(raw * 60.0 / max_possible, 2)
    
    # Таблица результатов
    doc.add_heading("Сводная таблица результатов", level=2)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Шкала"
    hdr[1].text = "Сырой балл"
    hdr[2].text = "Норм. (0..60)"
    hdr[3].text = "Интерпретация"
    
    # Заполняем таблицу
    if use_ai and ai_interpreter:
        # AI интерпретации
        for scale in sorted(scores_raw.keys()):
            row = table.add_row().cells
            row[0].text = scale
            row[1].text = str(scores_raw[scale])
            row[2].text = str(scaled_scores[scale])
            row[3].text = "См. подробный анализ ниже"
        
        # Добавляем подробный AI анализ
        doc.add_heading("Подробный анализ", level=2)
        
        if test_type.upper() == "PAEI":
            interpretation = ai_interpreter.interpret_paei(scores_raw, dialog_context)
        elif test_type.upper() == "DISC":
            interpretation = ai_interpreter.interpret_disc(scores_raw, dialog_context)
        elif test_type.upper() == "HEXACO":
            interpretation = ai_interpreter.interpret_hexaco(scores_raw, dialog_context)
        else:
            interpretation = "Тип теста не поддерживается для AI интерпретации"
        
        doc.add_paragraph(interpretation)
        
    else:
        # Статические интерпретации
        df_interp = pd.read_csv(interpretations_path)
        for scale in sorted(scores_raw.keys()):
            row = table.add_row().cells
            row[0].text = scale
            row[1].text = str(scores_raw[scale])
            row[2].text = str(scaled_scores[scale])
            row[3].text = _pick_interpretation(scale, scaled_scores[scale], df_interp)
    
    # Footnotes
    doc.add_paragraph()
    doc.add_paragraph("⚠️ Результаты носят ознакомительный характер и не заменяют профессиональную диагностику.")
    
    # Информация об использовании AI
    if use_ai and ai_interpreter:
        doc.add_paragraph()
        doc.add_paragraph(
            "🤖 Интерпретации результатов сгенерированы с помощью OpenAI GPT-3.5. "
            "Powered by OpenAI (https://openai.com)"
        )
    
    doc.save(str(out_path))
    return out_path


# Алиас для обратной совместимости
def render_report(*args, **kwargs):
    """Обёртка для обратной совместимости"""
    return render_enhanced_report(*args, **kwargs)