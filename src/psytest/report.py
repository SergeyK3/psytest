from pathlib import Path
from typing import Dict
from docx import Document
from docx.shared import Pt
import pandas as pd
from datetime import datetime

def _pick_interpretation(scale: str, scaled_score: float, df_interp: pd.DataFrame) -> str:
    pool = df_interp[df_interp["scale"] == scale]
    row = pool[(pool["range_low"] <= scaled_score) & (scaled_score <= pool["range_high"])]
    if not row.empty:
        return f"{row.iloc[0]['level']}: {row.iloc[0]['text']}"
    # если не попали в диапазон — просто вернём пусто
    return ""

def render_report(scores_raw: Dict[str, float],
                  items_df: pd.DataFrame,
                  interpretations_path: Path,
                  template_path: Path,
                  out_path: Path,
                  title: str = "Отчёт по психологическому тестированию (PAEI)",
                  participant: str = "Код участника: ______") -> Path:
    """
    scores_raw: {'P': raw, 'A': raw, ...}
    items_df: dataframe с вопросами (item_id, scale, reverse, ...)
    Нормируем сырые баллы к шкале 0..60, чтобы соотнести с интерпретациями.
    """
    doc = Document(str(template_path))

    # базовая типографика
    styles = doc.styles
    styles['Normal'].font.name = 'Times New Roman'
    styles['Normal'].font.size = Pt(12)

    doc.add_heading(title, level=1)
    doc.add_paragraph(f"{participant}")
    doc.add_paragraph(f"Дата: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    # Подготовим нормировку к 0..60, учитывая число пунктов по каждой шкале
    items_per_scale = items_df.groupby("scale")["item_id"].count().to_dict()
    max_per_scale = {k: v * 5 for k, v in items_per_scale.items()}  # 5 — максимум по слайдеру
    scaled_scores = {}
    for s, raw in scores_raw.items():
        max_possible = max_per_scale.get(s, max(raw, 1))  # защита от деления на ноль
        scaled_scores[s] = round(raw * 60.0 / max_possible, 2)

    # Загружаем интерпретации
    df_interp = pd.read_csv(interpretations_path)

    doc.add_heading("Сводная таблица результатов", level=2)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Шкала"
    hdr[1].text = "Сырой балл"
    hdr[2].text = "Норм. (0..60)"
    hdr[3].text = "Интерпретация"

    for scale in sorted(scores_raw.keys()):
        row = table.add_row().cells
        row[0].text = scale
        row[1].text = str(scores_raw[scale])
        row[2].text = str(scaled_scores[scale])
        row[3].text = _pick_interpretation(scale, scaled_scores[scale], df_interp)

    doc.add_paragraph()
    doc.add_paragraph("⚠️ Результаты носят ознакомительный характер и не заменяют профессиональную диагностику.")
    
    # Добавляем информацию об использовании AI (если используется)
    if hasattr(render_report, '_uses_ai') and render_report._uses_ai:
        doc.add_paragraph()
        doc.add_paragraph(
            "🤖 Интерпретации результатов сгенерированы с помощью OpenAI GPT-3.5. "
            "Powered by OpenAI (https://openai.com)"
        )

    doc.save(str(out_path))
    return out_path
